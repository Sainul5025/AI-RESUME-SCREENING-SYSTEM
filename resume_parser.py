"""
Resume Parser Module
--------------------
Extracts and cleans plain text from PDF and DOCX resume files.

This module is designed to be imported into app.py:

    from resume_parser import ResumeParser

    parser = ResumeParser()
    text = parser.extract_text("uploads/sample_resume.pdf")
"""

import logging
import os
import re
from pathlib import Path
from typing import Final, Set

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError

# ---------------------------------------------------------------------------
# Logging configuration
# ---------------------------------------------------------------------------

# Create a module-level logger for debugging parser activity.
logger = logging.getLogger(__name__)

# Supported resume file extensions (lowercase, with leading dot).
SUPPORTED_EXTENSIONS: Final[Set[str]] = {".pdf", ".docx"}


class UnsupportedFileTypeError(ValueError):
    """Raised when the file extension is not PDF or DOCX."""


class EmptyOrCorruptedFileError(ValueError):
    """Raised when a file is empty or cannot be read properly."""


class ResumeParser:
    """
    Parser for extracting readable text from resume files.

    Supports:
        - PDF files (.pdf) using PyPDF2
        - Word documents (.docx) using python-docx
    """

    def __init__(self) -> None:
        """Initialize the ResumeParser instance."""
        logger.debug("ResumeParser initialized.")

    def extract_text(self, file_path: str) -> str:
        """
        Extract and clean text from a resume file.

        This is the main entry point. It validates the file, detects the
        file type from the extension, and delegates to the correct extractor.

        Args:
            file_path: Absolute or relative path to the resume file.

        Returns:
            Cleaned plain-text content of the resume.

        Raises:
            FileNotFoundError: If the file does not exist.
            UnsupportedFileTypeError: If the extension is not .pdf or .docx.
            EmptyOrCorruptedFileError: If the file is empty or unreadable.
        """
        logger.info("Starting text extraction for: %s", file_path)

        # Step 1: Validate that the file exists and has a supported extension.
        validated_path = self._validate_file(file_path)

        # Step 2: Route to the correct extractor based on file extension.
        extension = validated_path.suffix.lower()

        if extension == ".pdf":
            raw_text = self.extract_pdf(str(validated_path))
        elif extension == ".docx":
            raw_text = self.extract_docx(str(validated_path))
        else:
            # This should not happen if _validate_file worked correctly.
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{extension}'. "
                f"Allowed types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        # Step 3: Clean the extracted text before returning it.
        cleaned_text = self._clean_text(raw_text)

        if not cleaned_text.strip():
            logger.error("Extracted text is empty for file: %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"The file '{file_path}' contains no readable text."
            )

        logger.info(
            "Successfully extracted %d characters from: %s",
            len(cleaned_text),
            file_path,
        )
        return cleaned_text

    def extract_pdf(self, file_path: str) -> str:
        """
        Extract raw text from a PDF resume using PyPDF2.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Raw text extracted from all pages of the PDF.

        Raises:
            FileNotFoundError: If the PDF file does not exist.
            EmptyOrCorruptedFileError: If the PDF is empty or corrupted.
        """
        logger.debug("Extracting text from PDF: %s", file_path)

        path = Path(file_path)

        if not path.is_file():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        try:
            with path.open("rb") as pdf_file:
                reader = PdfReader(pdf_file)

                # A PDF with zero pages is treated as corrupted/empty.
                if len(reader.pages) == 0:
                    raise EmptyOrCorruptedFileError(
                        f"The PDF file '{file_path}' has no pages."
                    )

                # Read text from each page and join with newlines.
                page_texts: list[str] = []
                for page_number, page in enumerate(reader.pages, start=1):
                    page_text = page.extract_text() or ""
                    logger.debug(
                        "Extracted %d characters from page %d.",
                        len(page_text),
                        page_number,
                    )
                    page_texts.append(page_text)

                extracted_text = "\n".join(page_texts)

        except PdfReadError as error:
            logger.exception("Failed to read PDF file: %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"The PDF file '{file_path}' is corrupted or invalid."
            ) from error
        except OSError as error:
            logger.exception("OS error while reading PDF: %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"Could not read the PDF file '{file_path}'."
            ) from error

        if not extracted_text.strip():
            raise EmptyOrCorruptedFileError(
                f"The PDF file '{file_path}' contains no extractable text."
            )

        return extracted_text

    def extract_docx(self, file_path: str) -> str:
        """
        Extract raw text from a DOCX resume using python-docx.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Raw text extracted from paragraphs and tables in the document.

        Raises:
            FileNotFoundError: If the DOCX file does not exist.
            EmptyOrCorruptedFileError: If the DOCX is empty or corrupted.
        """
        logger.debug("Extracting text from DOCX: %s", file_path)

        path = Path(file_path)

        if not path.is_file():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            document = Document(str(path))

            # Collect text from all paragraphs in the document body.
            paragraph_texts = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

            # Also collect text from tables (common in resume templates).
            table_texts: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    row_cells = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_cells:
                        table_texts.append(" | ".join(row_cells))

            extracted_text = "\n".join(paragraph_texts + table_texts)

        except PackageNotFoundError as error:
            logger.exception("Invalid DOCX package: %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"The DOCX file '{file_path}' is corrupted or invalid."
            ) from error
        except OSError as error:
            logger.exception("OS error while reading DOCX: %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"Could not read the DOCX file '{file_path}'."
            ) from error

        if not extracted_text.strip():
            raise EmptyOrCorruptedFileError(
                f"The DOCX file '{file_path}' contains no extractable text."
            )

        return extracted_text

    def _validate_file(self, file_path: str) -> Path:
        """
        Validate that the resume file exists and has a supported extension.

        Args:
            file_path: Path to the resume file.

        Returns:
            A resolved Path object for the validated file.

        Raises:
            FileNotFoundError: If the file does not exist.
            UnsupportedFileTypeError: If the extension is not allowed.
            EmptyOrCorruptedFileError: If the file size is zero bytes.
        """
        if not file_path or not str(file_path).strip():
            raise FileNotFoundError("No file path was provided.")

        path = Path(file_path).expanduser().resolve()

        if not path.exists():
            logger.error("File not found: %s", file_path)
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.is_file():
            logger.error("Path is not a file: %s", file_path)
            raise FileNotFoundError(f"Path is not a file: {file_path}")

        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            logger.error("Unsupported file extension: %s", extension)
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{extension}'. "
                f"Allowed types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if path.stat().st_size == 0:
            logger.error("File is empty (0 bytes): %s", file_path)
            raise EmptyOrCorruptedFileError(
                f"The file '{file_path}' is empty (0 bytes)."
            )

        logger.debug("File validation passed: %s", path)
        return path

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted resume text while keeping it readable.

        Cleaning steps:
            1. Normalize line endings to \\n
            2. Remove trailing spaces on each line
            3. Collapse multiple spaces/tabs within a line to a single space
            4. Collapse excessive blank lines (keep at most one empty line)

        Args:
            text: Raw text extracted from a resume file.

        Returns:
            Cleaned, readable plain text.
        """
        if not text:
            return ""

        # Normalize Windows/Mac line endings to Unix-style newlines.
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")

        cleaned_lines: list[str] = []
        previous_line_was_blank = False

        for line in normalized.split("\n"):
            # Remove leading/trailing whitespace from the line.
            stripped_line = line.strip()

            # Replace multiple spaces or tabs inside the line with one space.
            stripped_line = re.sub(r"[ \t]+", " ", stripped_line)

            if not stripped_line:
                # Keep only one blank line between content blocks.
                if not previous_line_was_blank:
                    cleaned_lines.append("")
                previous_line_was_blank = True
                continue

            cleaned_lines.append(stripped_line)
            previous_line_was_blank = False

        # Remove leading/trailing blank lines from the final output.
        cleaned_text = "\n".join(cleaned_lines).strip()

        logger.debug("Text cleaning complete. Final length: %d", len(cleaned_text))
        return cleaned_text


# ---------------------------------------------------------------------------
# Optional helper for easy import in app.py
# ---------------------------------------------------------------------------

def parse_resume(file_path: str) -> str:
    """
    Convenience function that creates a ResumeParser and extracts text.

    Args:
        file_path: Path to a PDF or DOCX resume file.

    Returns:
        Cleaned resume text as a string.
    """
    parser = ResumeParser()
    return parser.extract_text(file_path)


if __name__ == "__main__":
    # Example usage: test the parser on a sample resume file.
    #
    # 1. Place a sample file in the BACKEND folder, for example:
    #       sample_resume.pdf
    #    or
    #       sample_resume.docx
    #
    # 2. Run this file directly:
    #       python resume_parser.py

    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )

    # Change this path to your sample resume file for testing.
    sample_file = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "sample_resume.pdf",
    )

    parser = ResumeParser()

    try:
        resume_text = parser.extract_text(sample_file)
        print("=" * 60)
        print("RESUME TEXT EXTRACTION SUCCESSFUL")
        print("=" * 60)
        print(resume_text[:1000])  # Print first 1000 characters as preview
        if len(resume_text) > 1000:
            print("\n... (output truncated) ...")
        print("=" * 60)
        print(f"Total characters extracted: {len(resume_text)}")

    except FileNotFoundError as error:
        print(f"File error: {error}")
        print(
            "\nTip: Create or copy a sample PDF/DOCX file and update "
            "'sample_file' in the __main__ block."
        )
    except (UnsupportedFileTypeError, EmptyOrCorruptedFileError) as error:
        print(f"Parser error: {error}")